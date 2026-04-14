from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_report():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('SoulStream: An AI-Powered Mood Mapping Music Recommender', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n" * 5)
    
    subtitle = doc.add_paragraph('A Mini-Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    doc.add_paragraph("\n" * 2)

    submission_info = doc.add_paragraph(f"Generated on: {datetime.date.today().strftime('%B %d, %Y')}")
    submission_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Abstract ---
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "SoulStream is an innovative music recommendation system that utilizes real-time facial emotion recognition to deliver "
        "personalized audio experiences. Traditional systems rely on user history or static playlists; however, SoulStream "
        "introduces a dynamic 'reactive' approach. By integrating Google Gemini 1.5 Flash Vision API, the system identifies "
        "nuanced emotional states (Happy, Sad, Angry, Neutral) from a camera feed and maps them to a curated Spotify search query. "
        "The system also features a persistent preference engine and a high-fidelity Streamlit interface designed with glassmorphism aesthetics."
    )

    # --- Introduction ---
    doc.add_heading('2. Introduction', level=1)
    doc.add_paragraph(
        "Music plays a vital role in regulating human emotions. While streaming platforms have vast libraries, finding the right "
        "song for a specific mood often requires manual effort. SoulStream aims to automate this discovery process by using "
        "computer vision. It addresses the limitation of text-only mood input by allowing users to use their camera, creating a "
        "more natural and immediate interaction between the user and their music player."
    )

    # --- System Architecture ---
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph(
        "The architecture follows a modular AI-as-a-Service (AIaaS) pattern. The flow involves three primary components:"
    )
    doc.add_paragraph("1. Front-end Layer: A Streamlit-based web application with Camera Input and Text Fallback mechanisms.", style='List Bullet')
    doc.add_paragraph("2. Intelligence Layer: Google Gemini 1.5 Flash API for vision-to-emotion conversion and preference analysis.", style='List Bullet')
    doc.add_paragraph("3. Integration Layer: Spotipy (Spotify Web API wrapper) for real-time track fetching and playback.", style='List Bullet')

    # --- Methodology ---
    doc.add_heading('4. Methodology', level=1)
    doc.add_heading('4.1 Emotion Detection', level=2)
    doc.add_paragraph(
        "The system captures a frame from the user's camera. This image is sent to the Gemini 1.5 Flash model with a specialized prompt requesting "
        "an emotional analysis. The model returns a structured JSON containing the detected emotion, confidence scores, and musical keywords."
    )
    doc.add_heading('4.2 Keyword Harmonization', level=2)
    doc.add_paragraph(
        "SoulStream combines the detected 'Image Emotion' with user-specified 'Persistent Preferences' (e.g., favorite artists like The Weeknd). "
        "This merged context ensures that the recommendations are not just based on mood, but also on the user's specific musical taste."
    )

    # --- Implementation Details ---
    doc.add_heading('5. Implementation', level=1)
    doc.add_paragraph(
        "The project is implemented in Python, chosen for its vast ecosystem of AI and web libraries. Key libraries include:"
    )
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tool/Library'
    hdr_cells[1].text = 'Purpose'
    
    data = [
        ('Streamlit', 'Web UI and Camera Input'),
        ('Google GenAI', 'Gemini Vision and Text Analysis'),
        ('Spotipy', 'Spotify API Integration'),
        ('Pillow (PIL)', 'Image handling and preparation'),
        ('Dotenv', 'Secure API Key Management')
    ]

    for tool, purpose in data:
        row_cells = table.add_row().cells
        row_cells[0].text = tool
        row_cells[1].text = purpose

    # --- Conclusion ---
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "SoulStream demonstrates the power of Generative AI in enhancing personal entertainment. By transforming a simple music player "
        "into an emotionally aware assistant, the project showcases the potential of multi-modal AI in everyday applications. Future iterations "
        "could include real-time pulse detection via camera to further refine the emotional mapping."
    )

    # --- Save ---
    filename = "SoulStream_Mini_Project_Report.docx"
    doc.save(filename)
    print(f"Report successfully saved as {filename}")

if __name__ == "__main__":
    create_report()
